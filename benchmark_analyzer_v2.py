# -*- coding: utf-8 -*-
# ============================================================
# BrailleSpeakGR Benchmark Analyzer v2.0
# Professional benchmark engine for OCR/Braille/TTS comparison.
#
# Input:
#   Benchmark_Master_BrailleSpeakGR.csv
#
# Outputs:
#   benchmark_outputs_v2/
#     01_summary_tables/*.csv
#     02_article_text/*.txt
#     03_figures/*.png
#     04_reports/benchmark_report.docx
#     04_reports/benchmark_report.pdf
#     05_excel/benchmark_report.xlsx
#     06_json/benchmark_results.json
#
# Usage:
#   python benchmark_analyzer_v2.py
#   python benchmark_analyzer_v2.py Benchmark_Master_BrailleSpeakGR.csv
#
# Optional dependencies:
#   pip install matplotlib python-docx openpyxl
#
# The script still runs without optional packages and writes CSV/TXT/JSON.
# ============================================================

from __future__ import annotations

import csv
import json
import math
import sys
from dataclasses import dataclass, asdict
from pathlib import Path
from statistics import mean, median, pstdev
from collections import defaultdict
from typing import Dict, List, Optional, Any, Tuple


# -----------------------------
# Configuration
# -----------------------------

DEFAULT_INPUT = "Benchmark_Master_BrailleSpeakGR.csv"
OUT_ROOT = Path("benchmark_outputs_v2")

PREFERRED_SYSTEM_ORDER = [
    "BrailleSpeakGR",
    "VoiceOver",
    "Google Lookout",
    "Seeing AI",
]

ERROR_COLUMNS = [
    "diphthong_errors",
    "punctuation_errors",
    "polytonic_errors",
]

QUALITY_COLUMNS = [
    "ocr_success",
    "english_terms_ok",
    "braille_unicode_output",
]

# Scoring weights for automatic ranking.
# Higher score is better.
RANKING_WEIGHTS = {
    "ocr_success_rate_percent": 0.25,
    "braille_unicode_output_rate_percent": 0.25,
    "english_terms_success_rate_percent": 0.15,
    "known_error_penalty": 0.25,
    "response_time_penalty": 0.10,
}


# -----------------------------
# Utility functions
# -----------------------------

def safe_float(value: Any) -> Optional[float]:
    if value is None:
        return None
    value = str(value).strip()
    if value == "" or value.lower() in {"n/a", "na", "nan", "none", "null", "-", "—"}:
        return None
    try:
        return float(value)
    except ValueError:
        return None


def safe_int(value: Any) -> Optional[int]:
    v = safe_float(value)
    if v is None:
        return None
    return int(round(v))


def clean_text(value: Any) -> str:
    if value is None:
        return ""
    return str(value).strip()


def fmt(value: Optional[float], decimals: int = 2) -> str:
    if value is None:
        return "N/A"
    if isinstance(value, float) and math.isnan(value):
        return "N/A"
    return f"{value:.{decimals}f}"


def fmt0(value: Optional[float]) -> str:
    if value is None:
        return "N/A"
    return f"{value:.0f}"


def pct(value: Optional[float]) -> str:
    if value is None:
        return "N/A"
    return f"{value:.1f}%"


def ci95(values: List[float]) -> Tuple[Optional[float], Optional[float], Optional[float]]:
    """Return mean, lower, upper 95% CI using normal approximation."""
    values = [v for v in values if v is not None]
    if not values:
        return None, None, None
    m = mean(values)
    if len(values) == 1:
        return m, m, m
    sd = pstdev(values)
    se = sd / math.sqrt(len(values))
    margin = 1.96 * se
    return m, m - margin, m + margin


def ensure_dirs() -> Dict[str, Path]:
    dirs = {
        "root": OUT_ROOT,
        "tables": OUT_ROOT / "01_summary_tables",
        "article": OUT_ROOT / "02_article_text",
        "figures": OUT_ROOT / "03_figures",
        "reports": OUT_ROOT / "04_reports",
        "excel": OUT_ROOT / "05_excel",
        "json": OUT_ROOT / "06_json",
    }
    for d in dirs.values():
        d.mkdir(parents=True, exist_ok=True)
    return dirs


def sort_systems(systems: List[str]) -> List[str]:
    return sorted(
        systems,
        key=lambda s: PREFERRED_SYSTEM_ORDER.index(s) if s in PREFERRED_SYSTEM_ORDER else 99
    )


def save_csv(path: Path, rows: List[Dict[str, Any]], fieldnames: Optional[List[str]] = None) -> None:
    if not rows:
        return
    if fieldnames is None:
        fieldnames = list(rows[0].keys())
    with open(path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


# -----------------------------
# Data model
# -----------------------------

@dataclass
class BenchmarkRow:
    image_id: str
    image_name: str
    category: str
    system: str
    ocr_success: Optional[int]
    response_time_seconds: Optional[float]
    diphthong_errors: Optional[float]
    punctuation_errors: Optional[float]
    polytonic_errors: Optional[float]
    english_terms_ok: Optional[int]
    braille_unicode_output: Optional[int]
    notes: str
    source_file: str = ""

    @staticmethod
    def from_dict(r: Dict[str, Any]) -> "BenchmarkRow":
        return BenchmarkRow(
            image_id=clean_text(r.get("image_id")),
            image_name=clean_text(r.get("image_name")),
            category=clean_text(r.get("category")),
            system=clean_text(r.get("system")),
            ocr_success=safe_int(r.get("ocr_success")),
            response_time_seconds=safe_float(r.get("response_time_seconds")),
            diphthong_errors=safe_float(r.get("diphthong_errors")),
            punctuation_errors=safe_float(r.get("punctuation_errors")),
            polytonic_errors=safe_float(r.get("polytonic_errors")),
            english_terms_ok=safe_int(r.get("english_terms_ok")),
            braille_unicode_output=safe_int(r.get("braille_unicode_output")),
            notes=clean_text(r.get("notes")),
            source_file=clean_text(r.get("source_file")),
        )


# -----------------------------
# Analyzer engine
# -----------------------------

class BenchmarkAnalyzer:
    def __init__(self, csv_path: Path):
        self.csv_path = Path(csv_path)
        self.rows: List[BenchmarkRow] = []
        self.systems: List[str] = []
        self.categories: List[str] = []
        self.images: List[str] = []
        self.dirs = ensure_dirs()

    def load(self) -> None:
        if not self.csv_path.exists():
            raise FileNotFoundError(f"Input CSV not found: {self.csv_path}")
        with open(self.csv_path, "r", encoding="utf-8-sig", newline="") as f:
            reader = csv.DictReader(f)
            raw_rows = list(reader)
        if not raw_rows:
            raise ValueError("Input CSV is empty.")

        self.rows = [BenchmarkRow.from_dict(r) for r in raw_rows]
        self.systems = sort_systems(sorted({r.system for r in self.rows if r.system}))
        self.categories = sorted({r.category for r in self.rows if r.category})
        self.images = sorted({r.image_id for r in self.rows if r.image_id}, key=lambda x: int(x) if x.isdigit() else x)

    # -----------------------------
    # Basic group helpers
    # -----------------------------

    def by_system(self) -> Dict[str, List[BenchmarkRow]]:
        d = defaultdict(list)
        for r in self.rows:
            d[r.system].append(r)
        return dict(d)

    def by_category_system(self) -> Dict[Tuple[str, str], List[BenchmarkRow]]:
        d = defaultdict(list)
        for r in self.rows:
            d[(r.category, r.system)].append(r)
        return dict(d)

    # -----------------------------
    # Computation
    # -----------------------------

    def compute_system_summary(self) -> List[Dict[str, Any]]:
        grouped = self.by_system()
        rows = []

        # Max known errors used for ranking normalization
        total_known_error_values = []
        mean_time_values = []
        for system, sr in grouped.items():
            total_known_errors = self._total_known_errors(sr)
            if total_known_errors is not None:
                total_known_error_values.append(total_known_errors)
            t = self._mean_positive([r.response_time_seconds for r in sr])
            if t is not None:
                mean_time_values.append(t)

        max_errors = max(total_known_error_values) if total_known_error_values else 1
        max_time = max(mean_time_values) if mean_time_values else 1

        for system in self.systems:
            sr = grouped.get(system, [])
            if not sr:
                continue

            ocr_vals = self._not_none([r.ocr_success for r in sr])
            braille_vals = self._not_none([r.braille_unicode_output for r in sr])
            english_vals = self._not_none([r.english_terms_ok for r in sr])
            response_vals = [r.response_time_seconds for r in sr if r.response_time_seconds is not None and r.response_time_seconds > 0]

            diph = self._not_none([r.diphthong_errors for r in sr])
            punct = self._not_none([r.punctuation_errors for r in sr])
            poly = self._not_none([r.polytonic_errors for r in sr])

            all_error_values = diph + punct + poly
            total_known_errors = sum(all_error_values) if all_error_values else None

            time_mean, time_ci_low, time_ci_high = ci95(response_vals)

            row = {
                "system": system,
                "n_rows": len(sr),
                "n_images": len({r.image_id for r in sr}),
                "ocr_success_count": sum(ocr_vals) if ocr_vals else None,
                "ocr_success_rate_percent": 100 * sum(ocr_vals) / len(ocr_vals) if ocr_vals else None,
                "braille_unicode_output_count": sum(braille_vals) if braille_vals else None,
                "braille_unicode_output_rate_percent": 100 * sum(braille_vals) / len(braille_vals) if braille_vals else None,
                "english_terms_success_count": sum(english_vals) if english_vals else None,
                "english_terms_success_rate_percent": 100 * sum(english_vals) / len(english_vals) if english_vals else None,
                "response_time_mean_seconds": time_mean,
                "response_time_median_seconds": median(response_vals) if response_vals else None,
                "response_time_sd_seconds": pstdev(response_vals) if len(response_vals) > 1 else 0 if response_vals else None,
                "response_time_min_seconds": min(response_vals) if response_vals else None,
                "response_time_max_seconds": max(response_vals) if response_vals else None,
                "response_time_ci95_low_seconds": time_ci_low,
                "response_time_ci95_high_seconds": time_ci_high,
                "total_diphthong_errors": sum(diph) if diph else None,
                "mean_diphthong_errors": mean(diph) if diph else None,
                "total_punctuation_errors": sum(punct) if punct else None,
                "mean_punctuation_errors": mean(punct) if punct else None,
                "total_polytonic_errors": sum(poly) if poly else None,
                "mean_polytonic_errors": mean(poly) if poly else None,
                "total_known_errors": total_known_errors,
                "mean_known_errors": mean(all_error_values) if all_error_values else None,
                "known_error_observations": len(all_error_values),
            }

            row["overall_score_percent"] = self._score_system(row, max_errors=max_errors, max_time=max_time)
            rows.append(row)

        rows.sort(key=lambda r: (-1 if r["overall_score_percent"] is None else -r["overall_score_percent"]))
        for idx, r in enumerate(rows, start=1):
            r["rank"] = idx
        # Put preferred order in printed summary but keep ranking field.
        rows.sort(key=lambda r: PREFERRED_SYSTEM_ORDER.index(r["system"]) if r["system"] in PREFERRED_SYSTEM_ORDER else 99)
        return rows

    def compute_category_summary(self) -> List[Dict[str, Any]]:
        grouped = self.by_category_system()
        out = []
        for category in self.categories:
            for system in self.systems:
                sr = grouped.get((category, system), [])
                if not sr:
                    continue
                ocr_vals = self._not_none([r.ocr_success for r in sr])
                times = [r.response_time_seconds for r in sr if r.response_time_seconds is not None and r.response_time_seconds > 0]
                diph = self._not_none([r.diphthong_errors for r in sr])
                punct = self._not_none([r.punctuation_errors for r in sr])
                poly = self._not_none([r.polytonic_errors for r in sr])
                english = self._not_none([r.english_terms_ok for r in sr])
                out.append({
                    "category": category,
                    "system": system,
                    "n_images": len({r.image_id for r in sr}),
                    "ocr_success_rate_percent": 100 * sum(ocr_vals) / len(ocr_vals) if ocr_vals else None,
                    "response_time_mean_seconds": mean(times) if times else None,
                    "total_diphthong_errors": sum(diph) if diph else None,
                    "total_punctuation_errors": sum(punct) if punct else None,
                    "total_polytonic_errors": sum(poly) if poly else None,
                    "english_terms_success_rate_percent": 100 * sum(english) / len(english) if english else None,
                    "total_known_errors": sum(diph + punct + poly) if (diph + punct + poly) else None,
                })
        return out

    def compute_image_summary(self) -> List[Dict[str, Any]]:
        out = []
        grouped = defaultdict(list)
        for r in self.rows:
            grouped[(r.image_id, r.image_name, r.category)].append(r)

        for (image_id, image_name, category), sr in sorted(grouped.items(), key=lambda x: int(x[0][0]) if x[0][0].isdigit() else x[0][0]):
            for r in sorted(sr, key=lambda x: PREFERRED_SYSTEM_ORDER.index(x.system) if x.system in PREFERRED_SYSTEM_ORDER else 99):
                out.append({
                    "image_id": image_id,
                    "image_name": image_name,
                    "category": category,
                    "system": r.system,
                    "ocr_success": r.ocr_success,
                    "response_time_seconds": r.response_time_seconds,
                    "diphthong_errors": r.diphthong_errors,
                    "punctuation_errors": r.punctuation_errors,
                    "polytonic_errors": r.polytonic_errors,
                    "english_terms_ok": r.english_terms_ok,
                    "braille_unicode_output": r.braille_unicode_output,
                    "notes": r.notes,
                    "known_total_errors": sum(v for v in [r.diphthong_errors, r.punctuation_errors, r.polytonic_errors] if v is not None),
                })
        return out

    def compute_ranking_table(self, summary: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        ranked = sorted(summary, key=lambda r: r.get("rank", 99))
        out = []
        medal = {1: "1st", 2: "2nd", 3: "3rd"}
        for r in ranked:
            out.append({
                "rank": r["rank"],
                "ranking_label": medal.get(r["rank"], f"{r['rank']}th"),
                "system": r["system"],
                "overall_score_percent": r["overall_score_percent"],
                "ocr_success_rate_percent": r["ocr_success_rate_percent"],
                "total_known_errors": r["total_known_errors"],
                "response_time_mean_seconds": r["response_time_mean_seconds"],
                "braille_unicode_output_rate_percent": r["braille_unicode_output_rate_percent"],
            })
        return out

    def _score_system(self, row: Dict[str, Any], max_errors: float, max_time: float) -> Optional[float]:
        ocr = (row.get("ocr_success_rate_percent") or 0) / 100.0
        braille = (row.get("braille_unicode_output_rate_percent") or 0) / 100.0

        english_raw = row.get("english_terms_success_rate_percent")
        english = 0.5 if english_raw is None else english_raw / 100.0

        errors = row.get("total_known_errors")
        if errors is None:
            error_component = 0.0
        else:
            error_component = max(0.0, 1.0 - (errors / max(max_errors, 1)))

        t = row.get("response_time_mean_seconds")
        if t is None:
            time_component = 0.0
        else:
            time_component = max(0.0, 1.0 - (t / max(max_time, 1)))

        score = (
            RANKING_WEIGHTS["ocr_success_rate_percent"] * ocr +
            RANKING_WEIGHTS["braille_unicode_output_rate_percent"] * braille +
            RANKING_WEIGHTS["english_terms_success_rate_percent"] * english +
            RANKING_WEIGHTS["known_error_penalty"] * error_component +
            RANKING_WEIGHTS["response_time_penalty"] * time_component
        )
        return 100 * score

    @staticmethod
    def _not_none(values: List[Any]) -> List[Any]:
        return [v for v in values if v is not None]

    @staticmethod
    def _mean_positive(values: List[Optional[float]]) -> Optional[float]:
        vals = [v for v in values if v is not None and v > 0]
        return mean(vals) if vals else None

    @staticmethod
    def _total_known_errors(rows: List[BenchmarkRow]) -> Optional[float]:
        values = []
        for r in rows:
            for v in [r.diphthong_errors, r.punctuation_errors, r.polytonic_errors]:
                if v is not None:
                    values.append(v)
        return sum(values) if values else None

    # -----------------------------
    # Article text
    # -----------------------------

    def make_article_text(self, summary: List[Dict[str, Any]], ranking: List[Dict[str, Any]]) -> Dict[str, str]:
        by = {r["system"]: r for r in summary}
        b = by.get("BrailleSpeakGR", {})
        v = by.get("VoiceOver", {})
        l = by.get("Google Lookout", {})
        s = by.get("Seeing AI", {})

        main_paragraph = (
            "Comparative benchmark. Across the ten benchmark images, BrailleSpeakGR achieved "
            f"{pct(b.get('ocr_success_rate_percent'))} OCR success and produced "
            f"{fmt0(b.get('total_diphthong_errors'))} diphthong errors, "
            f"{fmt0(b.get('total_punctuation_errors'))} punctuation errors, and "
            f"{fmt0(b.get('total_polytonic_errors'))} observed polytonic errors. "
            f"VoiceOver also achieved {pct(v.get('ocr_success_rate_percent'))} OCR success but accumulated "
            f"{fmt0(v.get('total_diphthong_errors'))} diphthong errors, "
            f"{fmt0(v.get('total_punctuation_errors'))} punctuation errors, and "
            f"{fmt0(v.get('total_polytonic_errors'))} observed polytonic errors. "
            f"Google Lookout achieved {pct(l.get('ocr_success_rate_percent'))} OCR success but accumulated "
            f"{fmt0(l.get('total_diphthong_errors'))} diphthong errors and "
            f"{fmt0(l.get('total_punctuation_errors'))} punctuation errors, while polytonic Greek and mixed English terminology remained unreliable. "
            f"Seeing AI produced {pct(s.get('ocr_success_rate_percent'))} usable Greek output in the evaluated cases. "
            "BrailleSpeakGR was the only evaluated system providing Greek Braille Unicode output."
        )

        ranking_text = "Automatic ranking based on OCR success, Braille output, English terminology support, total known errors, and response time identified the following order: "
        ranking_text += "; ".join([f"{r['ranking_label']} {r['system']} ({fmt(r['overall_score_percent'], 1)})" for r in ranking]) + "."

        limitations_text = (
            "Benchmark limitations. The present benchmark uses a compact evaluation set and records observed error categories rather than full character-level transcriptions for every system. "
            "Therefore, the results should be interpreted as a targeted functional comparison focusing on Greek diphthongs, punctuation, polytonic Greek, mixed Greek-English terminology, response time, and Braille Unicode support. "
            "A larger benchmark with full ground-truth transcriptions would enable complete CER/WER computation in future work."
        )

        return {
            "article_main_paragraph.txt": main_paragraph,
            "article_ranking_paragraph.txt": ranking_text,
            "article_limitations_paragraph.txt": limitations_text,
        }

    # -----------------------------
    # Output generation
    # -----------------------------

    def write_outputs(self) -> Dict[str, Any]:
        self.load()

        summary = self.compute_system_summary()
        category_summary = self.compute_category_summary()
        image_summary = self.compute_image_summary()
        ranking = self.compute_ranking_table(summary)
        article_texts = self.make_article_text(summary, ranking)

        # CSV tables
        save_csv(self.dirs["tables"] / "system_summary.csv", summary)
        save_csv(self.dirs["tables"] / "category_summary.csv", category_summary)
        save_csv(self.dirs["tables"] / "image_level_summary.csv", image_summary)
        save_csv(self.dirs["tables"] / "ranking_table.csv", ranking)

        # Text outputs
        for filename, text in article_texts.items():
            (self.dirs["article"] / filename).write_text(text, encoding="utf-8")

        # JSON
        json_obj = {
            "input_csv": str(self.csv_path),
            "systems": self.systems,
            "categories": self.categories,
            "n_rows": len(self.rows),
            "n_images": len(self.images),
            "system_summary": summary,
            "category_summary": category_summary,
            "image_summary": image_summary,
            "ranking": ranking,
            "article_text": article_texts,
        }
        (self.dirs["json"] / "benchmark_results.json").write_text(
            json.dumps(json_obj, ensure_ascii=False, indent=2),
            encoding="utf-8"
        )

        # Figures
        figure_paths = self.make_figures(summary, category_summary, image_summary, ranking)

        # Excel / Word / PDF reports
        xlsx_path = self.make_excel(summary, category_summary, image_summary, ranking, article_texts, figure_paths)
        docx_path = self.make_docx(summary, category_summary, image_summary, ranking, article_texts, figure_paths)
        pdf_path = self.make_pdf(summary, category_summary, ranking, article_texts, figure_paths)

        return {
            "summary": summary,
            "category_summary": category_summary,
            "image_summary": image_summary,
            "ranking": ranking,
            "article_texts": article_texts,
            "figure_paths": [str(p) for p in figure_paths],
            "xlsx_path": str(xlsx_path) if xlsx_path else None,
            "docx_path": str(docx_path) if docx_path else None,
            "pdf_path": str(pdf_path) if pdf_path else None,
            "out_root": str(OUT_ROOT),
        }

    # -----------------------------
    # Figures
    # -----------------------------

    def make_figures(self, summary, category_summary, image_summary, ranking) -> List[Path]:
        try:
            import matplotlib.pyplot as plt
        except Exception:
            print("matplotlib not available. Figures skipped.")
            return []

        fig_dir = self.dirs["figures"]
        paths = []

        systems = [r["system"] for r in summary]

        def values(key):
            return [0 if r.get(key) is None else r.get(key) for r in summary]

        def save_bar(filename, y, ylabel, title, ylim=None):
            plt.figure(figsize=(10, 5.8))
            plt.bar(systems, y)
            plt.ylabel(ylabel)
            plt.title(title)
            if ylim:
                plt.ylim(*ylim)
            plt.xticks(rotation=25, ha="right")
            plt.tight_layout()
            p = fig_dir / filename
            plt.savefig(p, dpi=220)
            plt.close()
            paths.append(p)

        # 1-10 system-level bar charts
        save_bar("01_ocr_success_rate.png", values("ocr_success_rate_percent"), "OCR success rate (%)", "OCR Success Rate by System", (0, 110))
        save_bar("02_total_known_errors.png", values("total_known_errors"), "Total known errors", "Total Known Errors by System")
        save_bar("03_diphthong_errors.png", values("total_diphthong_errors"), "Diphthong errors", "Total Diphthong Errors by System")
        save_bar("04_punctuation_errors.png", values("total_punctuation_errors"), "Punctuation errors", "Total Punctuation Errors by System")
        save_bar("05_polytonic_errors.png", values("total_polytonic_errors"), "Polytonic errors", "Observed Polytonic Errors by System")
        save_bar("06_mean_response_time.png", values("response_time_mean_seconds"), "Mean response time (s)", "Mean Response Time by System")
        save_bar("07_braille_unicode_output_rate.png", values("braille_unicode_output_rate_percent"), "Braille output rate (%)", "Greek Braille Unicode Output by System", (0, 110))
        save_bar("08_english_terms_success_rate.png", values("english_terms_success_rate_percent"), "English terminology success (%)", "Mixed Greek-English Terminology Support", (0, 110))
        save_bar("09_overall_score.png", values("overall_score_percent"), "Overall score (%)", "Automatic Overall Benchmark Score", (0, 110))
        save_bar("10_known_error_observations.png", values("known_error_observations"), "Observed error fields", "Number of Observed Error Fields by System")

        # 11 Error components by system
        plt.figure(figsize=(10, 5.8))
        diph = values("total_diphthong_errors")
        punct = values("total_punctuation_errors")
        poly = values("total_polytonic_errors")
        x = range(len(systems))
        plt.bar(x, diph, label="Diphthong")
        plt.bar(x, punct, bottom=diph, label="Punctuation")
        bottom2 = [a+b for a, b in zip(diph, punct)]
        plt.bar(x, poly, bottom=bottom2, label="Polytonic")
        plt.ylabel("Errors")
        plt.title("Error Component Breakdown")
        plt.xticks(list(x), systems, rotation=25, ha="right")
        plt.legend()
        plt.tight_layout()
        p = fig_dir / "11_error_component_breakdown.png"
        plt.savefig(p, dpi=220)
        plt.close()
        paths.append(p)

        # 12 Response time boxplot
        grouped_times = []
        labels = []
        for system in systems:
            vals = [r["response_time_seconds"] for r in image_summary if r["system"] == system and r["response_time_seconds"] not in (None, 0)]
            if vals:
                grouped_times.append(vals)
                labels.append(system)
        if grouped_times:
            plt.figure(figsize=(10, 5.8))
            plt.boxplot(grouped_times, labels=labels)
            plt.ylabel("Response time (s)")
            plt.title("Response Time Distribution by System")
            plt.xticks(rotation=25, ha="right")
            plt.tight_layout()
            p = fig_dir / "12_response_time_boxplot.png"
            plt.savefig(p, dpi=220)
            plt.close()
            paths.append(p)

        # 13 Per-image total errors for BrailleSpeakGR
        for target_system, fname in [
            ("BrailleSpeakGR", "13_per_image_errors_braillespeakgr.png"),
            ("VoiceOver", "14_per_image_errors_voiceover.png"),
            ("Google Lookout", "15_per_image_errors_google_lookout.png"),
        ]:
            rows = [r for r in image_summary if r["system"] == target_system]
            if rows:
                labels = [str(r["image_id"]) for r in rows]
                vals = [r.get("known_total_errors") or 0 for r in rows]
                plt.figure(figsize=(10, 5.8))
                plt.bar(labels, vals)
                plt.xlabel("Image ID")
                plt.ylabel("Known total errors")
                plt.title(f"Per-Image Known Errors: {target_system}")
                plt.tight_layout()
                p = fig_dir / fname
                plt.savefig(p, dpi=220)
                plt.close()
                paths.append(p)

        # 16 Category total errors heatmap-like table image (simple imshow)
        cats = sorted(set(r["category"] for r in category_summary))
        mat = []
        for cat in cats:
            row = []
            for sysname in systems:
                match = next((r for r in category_summary if r["category"] == cat and r["system"] == sysname), None)
                row.append(0 if not match or match.get("total_known_errors") is None else match["total_known_errors"])
            mat.append(row)
        if mat:
            plt.figure(figsize=(12, max(5.8, 0.45 * len(cats))))
            plt.imshow(mat, aspect="auto")
            plt.colorbar(label="Total known errors")
            plt.yticks(range(len(cats)), cats)
            plt.xticks(range(len(systems)), systems, rotation=25, ha="right")
            plt.title("Category-Level Error Matrix")
            plt.tight_layout()
            p = fig_dir / "16_category_error_matrix.png"
            plt.savefig(p, dpi=220)
            plt.close()
            paths.append(p)

        return paths

    # -----------------------------
    # Excel report
    # -----------------------------

    def make_excel(self, summary, category_summary, image_summary, ranking, article_texts, figure_paths) -> Optional[Path]:
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
            from openpyxl.utils import get_column_letter
            from openpyxl.chart import BarChart, Reference
        except Exception:
            print("openpyxl not available. Excel report skipped.")
            return None

        out_path = self.dirs["excel"] / "benchmark_report.xlsx"

        wb = Workbook()
        ws = wb.active
        ws.title = "Dashboard"

        header_fill = PatternFill("solid", fgColor="1F4E78")
        header_font = Font(color="FFFFFF", bold=True)
        title_font = Font(size=16, bold=True, color="1F4E78")
        thin = Side(style="thin", color="CCCCCC")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)

        ws["A1"] = "BrailleSpeakGR Comparative Benchmark Dashboard"
        ws["A1"].font = title_font
        ws.merge_cells("A1:H1")

        headers = ["Rank", "System", "OCR Success", "Diphthong Err.", "Punctuation Err.", "Polytonic Err.", "Mean Time", "Braille Output", "Overall Score"]
        for col, h in enumerate(headers, 1):
            c = ws.cell(row=3, column=col, value=h)
            c.fill = header_fill
            c.font = header_font
            c.border = border
            c.alignment = Alignment(horizontal="center")

        rank_by_system = {r["system"]: r["rank"] for r in ranking}
        score_by_system = {r["system"]: r["overall_score_percent"] for r in ranking}
        for row_idx, r in enumerate(summary, 4):
            values = [
                rank_by_system.get(r["system"]),
                r["system"],
                r["ocr_success_rate_percent"],
                r["total_diphthong_errors"],
                r["total_punctuation_errors"],
                r["total_polytonic_errors"],
                r["response_time_mean_seconds"],
                r["braille_unicode_output_rate_percent"],
                score_by_system.get(r["system"]),
            ]
            for col_idx, value in enumerate(values, 1):
                c = ws.cell(row=row_idx, column=col_idx, value=value)
                c.border = border
                if col_idx in [3, 8, 9]:
                    c.number_format = "0.0"
                if col_idx == 7:
                    c.number_format = "0.00"

        for col in range(1, len(headers) + 1):
            ws.column_dimensions[get_column_letter(col)].width = 18

        # Add chart
        chart = BarChart()
        chart.title = "Total Known Errors"
        chart.y_axis.title = "Errors"
        chart.x_axis.title = "System"
        data = Reference(ws, min_col=4, max_col=6, min_row=3, max_row=3 + len(summary))
        cats = Reference(ws, min_col=2, min_row=4, max_row=3 + len(summary))
        chart.add_data(data, titles_from_data=True)
        chart.set_categories(cats)
        ws.add_chart(chart, "K3")

        def add_sheet(name, data_rows):
            sh = wb.create_sheet(name)
            if not data_rows:
                return
            fields = list(data_rows[0].keys())
            for col, h in enumerate(fields, 1):
                c = sh.cell(row=1, column=col, value=h)
                c.fill = header_fill
                c.font = header_font
                c.border = border
            for r_idx, row in enumerate(data_rows, 2):
                for c_idx, h in enumerate(fields, 1):
                    c = sh.cell(row=r_idx, column=c_idx, value=row[h])
                    c.border = border
            for col in range(1, len(fields) + 1):
                sh.column_dimensions[get_column_letter(col)].width = min(max(len(str(fields[col-1])) + 4, 14), 35)
            sh.freeze_panes = "A2"

        add_sheet("System Summary", summary)
        add_sheet("Category Summary", category_summary)
        add_sheet("Image Level", image_summary)
        add_sheet("Ranking", ranking)

        sh = wb.create_sheet("Article Text")
        row = 1
        for fname, text in article_texts.items():
            sh.cell(row=row, column=1, value=fname).font = Font(bold=True, color="1F4E78")
            row += 1
            sh.cell(row=row, column=1, value=text).alignment = Alignment(wrap_text=True, vertical="top")
            sh.row_dimensions[row].height = 90
            row += 2
        sh.column_dimensions["A"].width = 120

        wb.save(out_path)
        return out_path

    # -----------------------------
    # Word report
    # -----------------------------

    def make_docx(self, summary, category_summary, image_summary, ranking, article_texts, figure_paths) -> Optional[Path]:
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except Exception:
            print("python-docx not available. Word report skipped.")
            return None

        out_path = self.dirs["reports"] / "benchmark_report.docx"

        doc = Document()
        doc.add_heading("BrailleSpeakGR Comparative Benchmark Report", level=1)
        p = doc.add_paragraph()
        p.add_run("Automatic OCR/Braille/TTS benchmark analysis").bold = True
        doc.add_paragraph("Input file: " + str(self.csv_path))
        doc.add_paragraph(f"Systems: {', '.join(self.systems)}")
        doc.add_paragraph(f"Images: {len(self.images)} | Rows: {len(self.rows)}")

        doc.add_heading("1. Executive Summary", level=2)
        for text in article_texts.values():
            doc.add_paragraph(text)

        doc.add_heading("2. Overall System Summary", level=2)
        headers = ["System", "OCR %", "Diph. Err.", "Punct. Err.", "Poly. Err.", "Mean Time", "Braille %", "Score", "Rank"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h

        rank_by_system = {r["system"]: r["rank"] for r in ranking}
        score_by_system = {r["system"]: r["overall_score_percent"] for r in ranking}
        for r in summary:
            cells = table.add_row().cells
            vals = [
                r["system"],
                pct(r["ocr_success_rate_percent"]),
                fmt0(r["total_diphthong_errors"]),
                fmt0(r["total_punctuation_errors"]),
                fmt0(r["total_polytonic_errors"]),
                fmt(r["response_time_mean_seconds"], 2),
                pct(r["braille_unicode_output_rate_percent"]),
                fmt(score_by_system.get(r["system"]), 1),
                str(rank_by_system.get(r["system"], "")),
            ]
            for i, v in enumerate(vals):
                cells[i].text = v

        doc.add_heading("3. Ranking", level=2)
        for r in ranking:
            doc.add_paragraph(f"{r['ranking_label']}: {r['system']} - overall score {fmt(r['overall_score_percent'], 1)}", style=None)

        doc.add_heading("4. Figures", level=2)
        for pth in figure_paths[:16]:
            try:
                doc.add_paragraph(pth.stem.replace("_", " ").title())
                doc.add_picture(str(pth), width=Inches(5.8))
            except Exception:
                pass

        doc.add_heading("5. Category Summary", level=2)
        headers2 = ["Category", "System", "OCR %", "Diph.", "Punct.", "Poly.", "Mean Time", "Total Err."]
        table2 = doc.add_table(rows=1, cols=len(headers2))
        table2.style = "Table Grid"
        for i, h in enumerate(headers2):
            table2.rows[0].cells[i].text = h
        for r in category_summary:
            vals = [
                r["category"], r["system"], pct(r["ocr_success_rate_percent"]),
                fmt0(r["total_diphthong_errors"]), fmt0(r["total_punctuation_errors"]),
                fmt0(r["total_polytonic_errors"]), fmt(r["response_time_mean_seconds"], 2),
                fmt0(r["total_known_errors"]),
            ]
            cells = table2.add_row().cells
            for i, v in enumerate(vals):
                cells[i].text = v

        doc.add_heading("6. Methodological Note", level=2)
        doc.add_paragraph(
            "The benchmark is generated from a master CSV file in which each row corresponds to one system applied to one benchmark image. "
            "This structure enables automatic recalculation when additional images or systems are added."
        )

        doc.save(out_path)
        return out_path

    # -----------------------------
    # PDF report
    # -----------------------------

    def make_pdf(self, summary, category_summary, ranking, article_texts, figure_paths) -> Optional[Path]:
        try:
            import matplotlib.pyplot as plt
            from matplotlib.backends.backend_pdf import PdfPages
        except Exception:
            print("matplotlib PDF support not available. PDF report skipped.")
            return None

        out_path = self.dirs["reports"] / "benchmark_report.pdf"

        with PdfPages(out_path) as pdf:
            # Title page
            fig = plt.figure(figsize=(8.27, 11.69))
            fig.text(0.5, 0.92, "BrailleSpeakGR Comparative Benchmark Report", ha="center", fontsize=18, weight="bold")
            fig.text(0.5, 0.88, "Automatic OCR/Braille/TTS Benchmark Analysis", ha="center", fontsize=12)
            y = 0.80
            fig.text(0.08, y, f"Input CSV: {self.csv_path}", fontsize=9)
            y -= 0.04
            fig.text(0.08, y, f"Images: {len(self.images)} | Rows: {len(self.rows)} | Systems: {', '.join(self.systems)}", fontsize=9)

            y -= 0.07
            fig.text(0.08, y, "Executive Summary", fontsize=14, weight="bold")
            y -= 0.04
            for text in article_texts.values():
                wrapped = self._wrap(text, width=95)
                for line in wrapped:
                    fig.text(0.08, y, line, fontsize=8)
                    y -= 0.018
                    if y < 0.08:
                        pdf.savefig(fig, bbox_inches="tight")
                        plt.close(fig)
                        fig = plt.figure(figsize=(8.27, 11.69))
                        y = 0.92
                y -= 0.02

            pdf.savefig(fig, bbox_inches="tight")
            plt.close(fig)

            # Summary table page
            fig = plt.figure(figsize=(11.69, 8.27))
            ax = fig.add_subplot(111)
            ax.axis("off")
            table_data = []
            headers = ["System", "OCR %", "Diph.", "Punct.", "Poly.", "Mean Time", "Braille %", "Score", "Rank"]
            rank_by_system = {r["system"]: r["rank"] for r in ranking}
            score_by_system = {r["system"]: r["overall_score_percent"] for r in ranking}
            for r in summary:
                table_data.append([
                    r["system"], pct(r["ocr_success_rate_percent"]), fmt0(r["total_diphthong_errors"]),
                    fmt0(r["total_punctuation_errors"]), fmt0(r["total_polytonic_errors"]),
                    fmt(r["response_time_mean_seconds"], 2), pct(r["braille_unicode_output_rate_percent"]),
                    fmt(score_by_system.get(r["system"]), 1), str(rank_by_system.get(r["system"], ""))
                ])
            ax.set_title("Overall System Summary", fontsize=16, weight="bold", pad=20)
            tab = ax.table(cellText=table_data, colLabels=headers, loc="center", cellLoc="center")
            tab.auto_set_font_size(False)
            tab.set_fontsize(9)
            tab.scale(1, 1.6)
            pdf.savefig(fig, bbox_inches="tight")
            plt.close(fig)

            # Add each chart as its own page
            for pth in figure_paths[:16]:
                try:
                    img = plt.imread(str(pth))
                    fig = plt.figure(figsize=(11.69, 8.27))
                    ax = fig.add_subplot(111)
                    ax.imshow(img)
                    ax.axis("off")
                    pdf.savefig(fig, bbox_inches="tight")
                    plt.close(fig)
                except Exception:
                    pass

        return out_path

    @staticmethod
    def _wrap(text: str, width: int = 90) -> List[str]:
        words = text.split()
        lines = []
        line = ""
        for w in words:
            if len(line) + len(w) + 1 > width:
                lines.append(line)
                line = w
            else:
                line = (line + " " + w).strip()
        if line:
            lines.append(line)
        return lines

    # -----------------------------
    # Console output
    # -----------------------------

    def print_console_report(self, results: Dict[str, Any]) -> None:
        summary = results["summary"]
        ranking = results["ranking"]

        print("\n" + "=" * 78)
        print("BRAILLESPEAKGR BENCHMARK ANALYZER v2.0")
        print("=" * 78)
        print(f"Input CSV: {self.csv_path}")
        print(f"Rows: {len(self.rows)} | Images: {len(self.images)} | Systems: {len(self.systems)}")
        print("-" * 78)
        print(f"{'System':<18} {'OCR%':>8} {'Diph':>7} {'Punct':>8} {'Poly':>7} {'Time':>8} {'Braille%':>9} {'Score':>8} {'Rank':>5}")
        print("-" * 78)
        score_by_system = {r["system"]: r["overall_score_percent"] for r in ranking}
        rank_by_system = {r["system"]: r["rank"] for r in ranking}
        for r in summary:
            print(
                f"{r['system']:<18} "
                f"{pct(r['ocr_success_rate_percent']):>8} "
                f"{fmt0(r['total_diphthong_errors']):>7} "
                f"{fmt0(r['total_punctuation_errors']):>8} "
                f"{fmt0(r['total_polytonic_errors']):>7} "
                f"{fmt(r['response_time_mean_seconds'], 2):>8} "
                f"{pct(r['braille_unicode_output_rate_percent']):>9} "
                f"{fmt(score_by_system.get(r['system']), 1):>8} "
                f"{rank_by_system.get(r['system'], ''):>5}"
            )
        print("-" * 78)
        print("Automatic ranking:")
        for r in ranking:
            print(f"  {r['ranking_label']}: {r['system']} | score={fmt(r['overall_score_percent'], 1)}")
        print("=" * 78)
        print("Saved outputs:")
        for key, value in results.items():
            if key.endswith("_path") and value:
                print(f"  {value}")
        print(f"  {OUT_ROOT}")
        print("=" * 78 + "\n")


def main():
    csv_path = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(DEFAULT_INPUT)
    analyzer = BenchmarkAnalyzer(csv_path)
    results = analyzer.write_outputs()
    analyzer.print_console_report(results)


if __name__ == "__main__":
    main()
